#!/usr/bin/env python3
"""Convert a markdown resume to .docx using python-docx.

Usage:
    python3 make_docx.py <input.md> <output.docx>

Handles:
  - # H1  -> centered name (16pt bold)
  - ## H2 -> section header (11pt bold, uppercase)
  - ### H3 -> role/title line (10pt bold)
  - - bullet -> List Bullet style
  - **bold** inline formatting
  - [text](url) -> clickable hyperlinks (blue, underlined)
  - --- horizontal rule
  - plain paragraphs

For PDF output, convert the .docx with:
    libreoffice --headless --convert-to pdf <output.docx> --outdir <dir>
"""
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print(
        "ERROR: python-docx not installed. Run: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a clickable hyperlink to a paragraph (blue, underlined)."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str) -> None:
    """Add text with **bold** and [link](url) inline formatting."""
    parts = re.split(r"(\*\*.*?\*\*|\[.*?\]\(.*?\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[(.*?)\]\((.*?)\)", part)
            if m:
                add_hyperlink(paragraph, m.group(1), m.group(2))
            else:
                run = paragraph.add_run(part)

        else:
            run = paragraph.add_run(part)


def md_to_docx(md_path: str, docx_path: str) -> None:
    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    doc = Document()

    # Tight margins for a 1-page resume
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Compact font for 1-page fit
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # H1 - name (centered, large)
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(15)

        # H2 - section header
        elif line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(10)

        # H3 - role/title line
        elif line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, text)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

        # Horizontal rule
        elif line.strip() in ("---", "***"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("_" * 90)
            run.font.size = Pt(5)

        # Bullet point -> plain paragraph (ATS-friendly, no bullet character)
        elif line.lstrip().startswith("- "):
            text = line.lstrip()[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.15)
            add_inline(p, text)

        # Empty line - just spacing
        elif line.strip() == "":
            pass

        # Regular paragraph (contact line, summary text, etc.)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, line.strip())

        i += 1

    doc.save(str(docx_path))
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
